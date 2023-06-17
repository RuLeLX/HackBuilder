from tkinter import *
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
from random import randint

#Условия задач
quests = {
    1: ["1. Медиана вариационного ряда 12, 12, 16, 18, 18, 18, 19, 20, 21, 22, 22, 23, 25, 25 равна:\n	а) 19,5;	б) 20;	в) 19;	г) 21.",
        "1. Медиана вариационного ряда 12, 12, 16, 18, 18, 18, 20, 20, 21, 22, 22, 23, 25, 25 равна:\n	а) 19,5;	б) 20;	в) 19;	г) 21.",
        "1. Медиана вариационного ряда 12, 12, 16, 18, 18, 18, 19, 19, 21, 22, 22, 23, 25, 25 равна:\n	а) 19,5;	б) 20;	в) 19;	г) 21.",
        "1. Медиана вариационного ряда 12, 12, 16, 18, 18, 18, 21, 21, 22, 23, 24, 25, 25, 25 равна:\n	а) 19,5;	б) 20;	в) 19;	г) 21."
        ],
    2: ["2. Мода вариационного ряда 1, 5, 5, 5, 5, 9, 9, 9, 12 равна:\n  а) 5;	б) 1;	в) 9;	г) 12.",
        "2. Мода вариационного ряда 1, 1, 1, 1, 7, 9, 9, 9, 12 равна:\n  а) 5;	б) 1;	в) 9;	г) 12.",
        "2. Мода вариационного ряда 1, 4, 5, 5, 7, 9, 9, 9, 12 равна:\n  а) 5;	б) 1;	в) 9;	г) 12.",
        "2. Мода вариационного ряда 1, 4, 5, 5, 7, 9, 12, 12, 12 равна:\n   а) 5;	б) 1;	в) 9;	г) 12."
        ],
    3: ["3. Из генеральной совокупности извлечена выборка объема n = 72:",
        "3. Из генеральной совокупности извлечена выборка объема n = 68:",
        "3. Из генеральной совокупности извлечена выборка объема n = 62:",
        "3. Из генеральной совокупности извлечена выборка объема n = 80:",
        ],
    4: ["4. Из генеральной совокупности извлечена выборка объема n=100"
    ]

}
#Дополнение к условиям задач на случай, если нужно вставить текст после слайда/таблицы
progon = {
    3: "Тогда значение n3 равно:\n  а) 20;  б) 16; в) 10;   г) 28",
    4: "Тогда относительная частота варианты xi = 5 равна:\nа) 0,5;	б) 0,25;	в) 0,18;	г) 0,13."

}
#Данные для таблиц, которые указаны в качестве вариантов ответа
progon_table = {
    6:[ [['xi – xi+1', '0–4', '4–8', '8–12', '12–16'], ['ni', '20', '50', '120', '60']],
    [['xi – xi+1', '0–4', '4–8', '8–12', '12–16'], ['ni', '20', '50', '60', '120']],
    [['xi – xi+1', '0–4', '4–8', '8–12', '12–16'], ['ni', '120', '50', '20', '60']], 
    [['xi – xi+1', '0–4', '4–8', '8–12', '12–16'], ['ni', '48', '64', '20', '120']]],
}

#Картинки для некоторых заданий
picture = {
    
}
datatables = {
    3: [[['xi – xi+1', '0-2', '24', '46', '68', '810'], ['ni', '6', '14', 'n3', '20', '12']]],
    4: [[['xi', '3', '4', '5', '6', '7'], ['ni', '15', '20', 'n3', '8', '7']],
        [['xi', '3', '4', '5', '6', '7'], ['ni', '22', '30', 'n3', '10', '13']],
        [['xi', '3', '4', '5', '6', '7'], ['ni', '15', '35', 'n3', '25', '7']],
        [['xi', '3', '4', '5', '6', '7'], ['ni', '15', '40', 'n3', '25', '7']]
        ],  
    7: [[['Хi', '12', '14', '15', '19'], ['ni', '22', '14', '3', '1']]]
    }
result_table = []

class Generation:
    #функция которая переводит индекс вопроса в списке в букву варианта
    def int_litter(x):
        if x == 0:
            return 'а'
        elif x == 1:
            return 'б'
        elif x == 2:
            return 'в'
        elif x == 3:
            return 'г'
        
        
   #count - количество вариантов
    def command(count):
        test = Document()
        test.styles['Normal'].font.size = Pt(16)
        #цикл генерации вариантов
        for i in range(count):

            test.add_paragraph().add_run(' ' * 50 + f'Тест 2. Вариант {i+1}').font.size = Pt(16)
            test.add_paragraph().add_run('Фамилия______________________Группа__________').font.size = Pt(16)

            variant = [] #список вариантов ответа генерируемого варианта, который будет помещён в конечный список ответов всех вариантов
            for quest in quests.items(): #items - получение пары ключ/значение
                change_mode = "string"  #в каких-то случаях варьирование происходит за счёт изменения строки условия задачи,а где-то изменяются данные таблицы/диаграммы
                if quest[0] >= 4 and quest[0] <= 7:
                    change_mode = "another"


                rand_num = randint(0, len(quest[1])-1)
                q = quest[1][rand_num]
                if change_mode == 'string':
                    variant.append(Generation.int_litter(rand_num))
                par = test.add_paragraph(q)

                #таблицы/прогоны/картинки есть не для всех заданий, потому избегаем потенциальной ошибки
                try:
                    rand_num = randint(0, len(picture[quest[0]])-1)
                    test.add_picture(picture[quest[0]][rand_num])
                    if change_mode == "another":
                        variant.append(Generation.int_litter(rand_num))
                except KeyError:
                    pass
                
                #таблицы, которые следуют за нумерованным параграфом
                try:
                    

                    rand_num = randint(0, len(datatables[quest[0]])-1)
                    table = test.add_table(rows=len(datatables[quest[0]][rand_num]), cols=len(datatables[quest[0]][rand_num][0]))
                    table.style = 'Table Grid'
                    #заполнение таблицы
                    for i in range(len(datatables[quest[0]][rand_num])):
                        for j in range(len(datatables[quest[0]][rand_num][0])):
                            cell = table.cell(i, j)
                            cell.text = datatables[quest[0]][rand_num][i][j]
                    if change_mode == "another":
                        variant.append(Generation.int_litter(rand_num))
                except KeyError:
                    pass

                try:
                    test.add_paragraph(progon[quest[0]])
                except KeyError:
                    continue                
                
                



            result_table.append(variant)   

        test.save('test.docx')
        print(result_table)
    
    def __init__(self, root, text, x, y, command):
        self.root = root
        self.text = text
        self.x = x
        self.y = y
        self.command = command
        Button(root, text = text, command=command).place(x=x, y=y)
    
        
    
        